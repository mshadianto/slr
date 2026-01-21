"""
BiblioAgent AI - Narrative Generator Module
============================================
Generates formal Indonesian academic narrative for "Results and Discussion"
chapter from SLR results data.

Features:
- PRISMA flow diagram narrative
- Study characteristics summary
- Quality assessment narrative
- Thematic synthesis discussion
- Export to Markdown/Word
"""

import logging
from typing import Dict, List, Optional, Any
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum

logger = logging.getLogger(__name__)


class NarrativeSection(Enum):
    """Sections of Results and Discussion chapter."""
    PRISMA_FLOW = "prisma_flow"
    STUDY_CHARACTERISTICS = "study_characteristics"
    QUALITY_ASSESSMENT = "quality_assessment"
    THEMATIC_SYNTHESIS = "thematic_synthesis"
    DISCUSSION = "discussion"
    LIMITATIONS = "limitations"


@dataclass
class NarrativeConfig:
    """Configuration for narrative generation."""
    language: str = "id"  # id = Indonesian, en = English
    formality_level: str = "formal"  # formal, semi-formal
    citation_style: str = "apa"  # apa, vancouver, harvard
    include_statistics: bool = True
    include_tables: bool = True
    max_section_length: int = 1500  # words per section
    target_audience: str = "academic"  # academic, general


@dataclass
class GeneratedNarrative:
    """Result of narrative generation."""
    section: NarrativeSection
    title: str
    content: str
    word_count: int
    tables: List[str] = field(default_factory=list)
    figures: List[str] = field(default_factory=list)
    citations: List[str] = field(default_factory=list)
    generated_at: str = field(default_factory=lambda: datetime.now().isoformat())


class NarrativeGenerator:
    """
    Generates formal Indonesian academic narrative for SLR results.

    Sections generated:
    1. PRISMA Flow - Selection process narrative
    2. Study Characteristics - Summary of included studies
    3. Quality Assessment - JBI quality scores discussion
    4. Thematic Synthesis - Key themes and findings
    5. Discussion - Interpretation and implications
    6. Limitations - Study limitations
    """

    # Indonesian academic phrases
    INDONESIAN_PHRASES = {
        'intro': [
            "Berdasarkan hasil pencarian literatur",
            "Dari proses seleksi yang dilakukan",
            "Hasil tinjauan sistematis menunjukkan",
            "Analisis terhadap literatur yang diperoleh",
        ],
        'transition': [
            "Selanjutnya",
            "Lebih lanjut",
            "Di samping itu",
            "Selain itu",
            "Adapun",
            "Berkenaan dengan hal tersebut",
        ],
        'conclusion': [
            "Dengan demikian",
            "Berdasarkan temuan di atas",
            "Dapat disimpulkan bahwa",
            "Hasil analisis menunjukkan",
        ],
        'comparison': [
            "Dibandingkan dengan",
            "Sejalan dengan",
            "Berbeda dengan",
            "Konsisten dengan",
        ],
    }

    def __init__(
        self,
        anthropic_client=None,
        config: NarrativeConfig = None
    ):
        """
        Initialize Narrative Generator.

        Args:
            anthropic_client: Anthropic client for LLM generation
            config: Narrative configuration
        """
        self.anthropic_client = anthropic_client
        self.config = config or NarrativeConfig()
        self.generated_sections: Dict[NarrativeSection, GeneratedNarrative] = {}

    async def generate_full_chapter(
        self,
        slr_results: Dict[str, Any]
    ) -> Dict[str, GeneratedNarrative]:
        """
        Generate complete Results and Discussion chapter.

        Args:
            slr_results: SLR results containing papers, stats, quality scores

        Returns:
            Dictionary of generated narrative sections
        """
        sections = {}

        # 1. PRISMA Flow Narrative
        logger.info("Generating PRISMA flow narrative...")
        sections[NarrativeSection.PRISMA_FLOW] = await self.generate_prisma_narrative(
            slr_results.get('prisma_stats', {}),
            slr_results.get('exclusion_reasons', {})
        )

        # 2. Study Characteristics
        logger.info("Generating study characteristics narrative...")
        sections[NarrativeSection.STUDY_CHARACTERISTICS] = await self.generate_characteristics_narrative(
            slr_results.get('acquired_papers', [])
        )

        # 3. Quality Assessment
        logger.info("Generating quality assessment narrative...")
        sections[NarrativeSection.QUALITY_ASSESSMENT] = await self.generate_quality_narrative(
            slr_results.get('quality_assessed_papers', [])
        )

        # 4. Thematic Synthesis
        logger.info("Generating thematic synthesis narrative...")
        sections[NarrativeSection.THEMATIC_SYNTHESIS] = await self.generate_thematic_narrative(
            slr_results.get('acquired_papers', []),
            slr_results.get('themes', [])
        )

        # 5. Discussion
        logger.info("Generating discussion narrative...")
        sections[NarrativeSection.DISCUSSION] = await self.generate_discussion_narrative(
            slr_results
        )

        # 6. Limitations
        logger.info("Generating limitations narrative...")
        sections[NarrativeSection.LIMITATIONS] = await self.generate_limitations_narrative(
            slr_results
        )

        self.generated_sections = sections
        return sections

    async def generate_prisma_narrative(
        self,
        prisma_stats: Dict[str, int],
        exclusion_reasons: Dict[str, int] = None
    ) -> GeneratedNarrative:
        """Generate PRISMA flow selection process narrative."""

        # Extract stats with defaults
        identified = prisma_stats.get('identified', 0)
        duplicates = prisma_stats.get('duplicates_removed', 0)
        screened = prisma_stats.get('screened', 0)
        excluded_screening = prisma_stats.get('excluded_screening', 0)
        sought = prisma_stats.get('sought_retrieval', 0)
        not_retrieved = prisma_stats.get('not_retrieved', 0)
        assessed = prisma_stats.get('assessed_eligibility', 0)
        excluded_eligibility = prisma_stats.get('excluded_eligibility', 0)
        included = prisma_stats.get('included_synthesis', 0)

        # Build narrative
        if self.anthropic_client:
            narrative = await self._generate_with_llm(
                section="PRISMA Flow",
                data={
                    'identified': identified,
                    'duplicates': duplicates,
                    'screened': screened,
                    'excluded_screening': excluded_screening,
                    'sought': sought,
                    'not_retrieved': not_retrieved,
                    'assessed': assessed,
                    'excluded_eligibility': excluded_eligibility,
                    'included': included,
                    'exclusion_reasons': exclusion_reasons or {}
                },
                template=self._get_prisma_template()
            )
        else:
            narrative = self._generate_prisma_fallback(
                identified, duplicates, screened, excluded_screening,
                sought, not_retrieved, assessed, excluded_eligibility,
                included, exclusion_reasons
            )

        return GeneratedNarrative(
            section=NarrativeSection.PRISMA_FLOW,
            title="4.1 Proses Seleksi Studi (PRISMA Flow)",
            content=narrative,
            word_count=len(narrative.split())
        )

    def _generate_prisma_fallback(
        self,
        identified: int,
        duplicates: int,
        screened: int,
        excluded_screening: int,
        sought: int,
        not_retrieved: int,
        assessed: int,
        excluded_eligibility: int,
        included: int,
        exclusion_reasons: Dict = None
    ) -> str:
        """Generate PRISMA narrative without LLM."""

        after_duplicates = identified - duplicates
        retrieved = sought - not_retrieved

        narrative = f"""Proses seleksi studi dalam tinjauan sistematis ini mengikuti pedoman PRISMA 2020 (Preferred Reporting Items for Systematic Reviews and Meta-Analyses). Pencarian literatur dilakukan secara komprehensif melalui basis data elektronik yang relevan.

Hasil pencarian awal mengidentifikasi sebanyak {identified:,} artikel dari berbagai sumber basis data. Setelah dilakukan penghapusan duplikat, diperoleh {after_duplicates:,} artikel yang unik untuk dilakukan skrining lebih lanjut. Proses deduplikasi menghapus {duplicates:,} artikel ({(duplicates/identified*100) if identified > 0 else 0:.1f}%) yang teridentifikasi sebagai duplikat.

Pada tahap skrining judul dan abstrak, sebanyak {screened:,} artikel dievaluasi berdasarkan kriteria inklusi dan eksklusi yang telah ditetapkan. Dari proses ini, {excluded_screening:,} artikel dieksklusi karena tidak memenuhi kriteria kelayakan. """

        if exclusion_reasons:
            narrative += "Alasan eksklusi meliputi: "
            reasons = [f"{reason} (n={count})" for reason, count in exclusion_reasons.items()]
            narrative += ", ".join(reasons) + ". "

        narrative += f"""

Selanjutnya, dilakukan pencarian teks lengkap (full-text) terhadap {sought:,} artikel yang lolos skrining awal. Dari jumlah tersebut, {not_retrieved:,} artikel tidak dapat diperoleh teks lengkapnya karena kendala akses atau ketersediaan. Sebanyak {retrieved:,} artikel berhasil diperoleh dan dinilai kelayakannya secara menyeluruh.

Pada tahap penilaian kelayakan akhir, {assessed:,} artikel dievaluasi secara mendalam. Sebanyak {excluded_eligibility:,} artikel dieksklusi pada tahap ini karena tidak memenuhi kriteria metodologis atau substansial yang dipersyaratkan.

Dengan demikian, sebanyak {included:,} artikel memenuhi seluruh kriteria dan diikutsertakan dalam sintesis kualitatif tinjauan sistematis ini. Diagram alur PRISMA yang menggambarkan proses seleksi studi secara lengkap disajikan pada Gambar 4.1."""

        return narrative

    async def generate_characteristics_narrative(
        self,
        papers: List[Dict]
    ) -> GeneratedNarrative:
        """Generate study characteristics summary narrative."""

        if not papers:
            return GeneratedNarrative(
                section=NarrativeSection.STUDY_CHARACTERISTICS,
                title="4.2 Karakteristik Studi yang Diinklusi",
                content="Tidak ada studi yang diinklusi dalam tinjauan sistematis ini.",
                word_count=10
            )

        # Analyze papers
        years = [p.get('year', 0) for p in papers if p.get('year')]
        year_range = f"{min(years)}-{max(years)}" if years else "N/A"

        # Count by year
        year_counts = {}
        for y in years:
            year_counts[y] = year_counts.get(y, 0) + 1

        # Count by venue/journal
        venues = {}
        for p in papers:
            venue = p.get('venue', 'Tidak diketahui')
            if venue:
                venues[venue] = venues.get(venue, 0) + 1

        # Count by method (if available)
        methods = {}
        for p in papers:
            method = p.get('study_design', p.get('method', 'Tidak disebutkan'))
            methods[method] = methods.get(method, 0) + 1

        # Countries (if available)
        countries = {}
        for p in papers:
            country = p.get('country', None)
            if country:
                countries[country] = countries.get(country, 0) + 1

        if self.anthropic_client:
            narrative = await self._generate_with_llm(
                section="Study Characteristics",
                data={
                    'total_papers': len(papers),
                    'year_range': year_range,
                    'year_distribution': year_counts,
                    'venues': venues,
                    'methods': methods,
                    'countries': countries,
                    'papers': papers[:10]  # Sample for context
                },
                template=self._get_characteristics_template()
            )
        else:
            narrative = self._generate_characteristics_fallback(
                len(papers), year_range, year_counts, venues, methods, countries
            )

        # Generate table
        table = self._generate_characteristics_table(papers)

        return GeneratedNarrative(
            section=NarrativeSection.STUDY_CHARACTERISTICS,
            title="4.2 Karakteristik Studi yang Diinklusi",
            content=narrative,
            word_count=len(narrative.split()),
            tables=[table]
        )

    def _generate_characteristics_fallback(
        self,
        total: int,
        year_range: str,
        year_counts: Dict,
        venues: Dict,
        methods: Dict,
        countries: Dict
    ) -> str:
        """Generate characteristics narrative without LLM."""

        narrative = f"""Sebanyak {total} studi diikutsertakan dalam tinjauan sistematis ini. Studi-studi tersebut dipublikasikan dalam rentang waktu {year_range}. """

        # Year distribution
        if year_counts:
            sorted_years = sorted(year_counts.items(), key=lambda x: x[1], reverse=True)
            peak_year = sorted_years[0]
            narrative += f"Publikasi terbanyak terjadi pada tahun {peak_year[0]} dengan {peak_year[1]} studi ({peak_year[1]/total*100:.1f}%). "

            recent_years = [y for y in year_counts.keys() if y >= 2020]
            if recent_years:
                recent_count = sum(year_counts[y] for y in recent_years)
                narrative += f"Sebanyak {recent_count} studi ({recent_count/total*100:.1f}%) dipublikasikan dalam lima tahun terakhir, menunjukkan peningkatan minat penelitian pada topik ini. "

        # Venues
        if venues:
            sorted_venues = sorted(venues.items(), key=lambda x: x[1], reverse=True)[:5]
            narrative += f"\n\nStudi-studi yang diinklusi dipublikasikan pada {len(venues)} jurnal atau venue yang berbeda. "
            if sorted_venues:
                top_venue = sorted_venues[0]
                narrative += f"Jurnal dengan publikasi terbanyak adalah {top_venue[0]} dengan {top_venue[1]} artikel. "

        # Methods
        if methods and len(methods) > 1:
            narrative += "\n\nDari segi desain penelitian, "
            method_desc = [f"{method} ({count} studi, {count/total*100:.1f}%)"
                          for method, count in sorted(methods.items(), key=lambda x: x[1], reverse=True)]
            narrative += ", ".join(method_desc[:5]) + ". "

        # Countries
        if countries:
            narrative += f"\n\nSecara geografis, studi-studi tersebut berasal dari {len(countries)} negara yang berbeda. "
            sorted_countries = sorted(countries.items(), key=lambda x: x[1], reverse=True)[:3]
            if sorted_countries:
                narrative += "Negara dengan kontribusi terbanyak adalah " + \
                    ", ".join([f"{c[0]} ({c[1]} studi)" for c in sorted_countries]) + ". "

        narrative += "\n\nRingkasan karakteristik studi yang diinklusi disajikan pada Tabel 4.1."

        return narrative

    def _generate_characteristics_table(self, papers: List[Dict]) -> str:
        """Generate Markdown table of study characteristics."""

        header = "| No | Penulis (Tahun) | Judul | Jurnal | Metode | Temuan Utama |"
        separator = "|:---:|:---|:---|:---|:---|:---|"
        rows = [header, separator]

        for i, paper in enumerate(papers[:20], 1):  # Limit to 20
            authors = paper.get('authors', ['N/A'])
            first_author = authors[0] if authors else 'N/A'
            if len(authors) > 1:
                first_author += " et al."

            year = paper.get('year', 'N/A')
            title = paper.get('title', 'N/A')[:60] + "..." if len(paper.get('title', '')) > 60 else paper.get('title', 'N/A')
            venue = paper.get('venue', 'N/A')[:30] if paper.get('venue') else 'N/A'
            method = paper.get('study_design', paper.get('method', 'N/A'))
            findings = paper.get('tldr', paper.get('key_findings', 'N/A'))
            if findings and len(findings) > 80:
                findings = findings[:80] + "..."

            row = f"| {i} | {first_author} ({year}) | {title} | {venue} | {method} | {findings} |"
            rows.append(row)

        return "\n".join(rows)

    async def generate_quality_narrative(
        self,
        papers: List[Dict]
    ) -> GeneratedNarrative:
        """Generate quality assessment narrative."""

        if not papers:
            return GeneratedNarrative(
                section=NarrativeSection.QUALITY_ASSESSMENT,
                title="4.3 Penilaian Kualitas Studi",
                content="Tidak ada studi yang dinilai kualitasnya.",
                word_count=8
            )

        # Analyze quality scores
        quality_scores = [p.get('quality_score', 0) for p in papers if 'quality_score' in p]

        if not quality_scores:
            quality_scores = [p.get('retrieval_quality_score', 0) for p in papers]

        # Categorize
        high = len([s for s in quality_scores if s >= 0.8])
        moderate = len([s for s in quality_scores if 0.6 <= s < 0.8])
        low = len([s for s in quality_scores if 0.4 <= s < 0.6])
        critical = len([s for s in quality_scores if s < 0.4])

        avg_score = sum(quality_scores) / len(quality_scores) if quality_scores else 0

        if self.anthropic_client:
            narrative = await self._generate_with_llm(
                section="Quality Assessment",
                data={
                    'total': len(papers),
                    'high': high,
                    'moderate': moderate,
                    'low': low,
                    'critical': critical,
                    'average_score': avg_score,
                    'papers': papers[:5]
                },
                template=self._get_quality_template()
            )
        else:
            narrative = self._generate_quality_fallback(
                len(papers), high, moderate, low, critical, avg_score
            )

        return GeneratedNarrative(
            section=NarrativeSection.QUALITY_ASSESSMENT,
            title="4.3 Penilaian Kualitas Studi",
            content=narrative,
            word_count=len(narrative.split())
        )

    def _generate_quality_fallback(
        self,
        total: int,
        high: int,
        moderate: int,
        low: int,
        critical: int,
        avg_score: float
    ) -> str:
        """Generate quality assessment narrative without LLM."""

        narrative = f"""Penilaian kualitas metodologis dilakukan terhadap {total} studi yang diinklusi menggunakan instrumen JBI Critical Appraisal Tools yang disesuaikan dengan desain masing-masing studi. Penilaian mencakup aspek-aspek seperti kejelasan tujuan penelitian, kesesuaian desain, validitas pengukuran, dan kelengkapan pelaporan hasil.

Hasil penilaian menunjukkan bahwa rata-rata skor kualitas studi adalah {avg_score*100:.1f} dari skala 100. """

        if high > 0:
            narrative += f"Sebanyak {high} studi ({high/total*100:.1f}%) dikategorikan memiliki kualitas tinggi (skor ≥80), "
        if moderate > 0:
            narrative += f"{moderate} studi ({moderate/total*100:.1f}%) berkualitas sedang (skor 60-79), "
        if low > 0:
            narrative += f"{low} studi ({low/total*100:.1f}%) berkualitas rendah (skor 40-59), "
        if critical > 0:
            narrative += f"dan {critical} studi ({critical/total*100:.1f}%) memiliki kualitas sangat rendah (skor <40). "

        # Interpretation
        if avg_score >= 0.7:
            narrative += "\n\nSecara keseluruhan, mayoritas studi yang diinklusi memiliki kualitas metodologis yang memadai, sehingga temuan dari tinjauan sistematis ini dapat dianggap cukup reliabel."
        elif avg_score >= 0.5:
            narrative += "\n\nKualitas metodologis studi yang diinklusi bervariasi, sehingga interpretasi temuan perlu dilakukan dengan mempertimbangkan keterbatasan tersebut."
        else:
            narrative += "\n\nKualitas metodologis studi yang diinklusi secara umum masih perlu ditingkatkan. Hal ini menjadi pertimbangan penting dalam menginterpretasikan temuan tinjauan sistematis ini."

        narrative += "\n\nRincian hasil penilaian kualitas masing-masing studi disajikan pada Tabel 4.2."

        return narrative

    async def generate_thematic_narrative(
        self,
        papers: List[Dict],
        themes: List[Dict] = None
    ) -> GeneratedNarrative:
        """Generate thematic synthesis narrative."""

        if not papers:
            return GeneratedNarrative(
                section=NarrativeSection.THEMATIC_SYNTHESIS,
                title="4.4 Sintesis Tematik",
                content="Tidak ada data untuk sintesis tematik.",
                word_count=7
            )

        # Extract themes from papers if not provided
        if not themes:
            themes = self._extract_themes_from_papers(papers)

        if self.anthropic_client:
            narrative = await self._generate_with_llm(
                section="Thematic Synthesis",
                data={
                    'papers': papers,
                    'themes': themes,
                    'total_papers': len(papers)
                },
                template=self._get_thematic_template()
            )
        else:
            narrative = self._generate_thematic_fallback(papers, themes)

        return GeneratedNarrative(
            section=NarrativeSection.THEMATIC_SYNTHESIS,
            title="4.4 Sintesis Tematik",
            content=narrative,
            word_count=len(narrative.split())
        )

    def _extract_themes_from_papers(self, papers: List[Dict]) -> List[Dict]:
        """Extract themes from papers based on keywords and TL;DR."""
        # Simple keyword-based theme extraction
        theme_keywords = {}

        for paper in papers:
            keywords = paper.get('keywords', [])
            if isinstance(keywords, str):
                keywords = [k.strip() for k in keywords.split(',')]

            for kw in keywords:
                if kw:
                    theme_keywords[kw.lower()] = theme_keywords.get(kw.lower(), 0) + 1

        # Sort by frequency
        sorted_themes = sorted(theme_keywords.items(), key=lambda x: x[1], reverse=True)[:10]

        return [{'name': t[0], 'count': t[1]} for t in sorted_themes]

    def _generate_thematic_fallback(
        self,
        papers: List[Dict],
        themes: List[Dict]
    ) -> str:
        """Generate thematic synthesis without LLM."""

        narrative = f"""Berdasarkan analisis terhadap {len(papers)} studi yang diinklusi, beberapa tema utama teridentifikasi dalam literatur. Sintesis tematik dilakukan dengan menggunakan pendekatan analisis konten untuk mengidentifikasi pola dan kategori yang muncul dari temuan studi-studi tersebut.

"""

        if themes:
            narrative += "Tema-tema utama yang teridentifikasi meliputi:\n\n"
            for i, theme in enumerate(themes[:5], 1):
                theme_name = theme.get('name', f'Tema {i}')
                theme_count = theme.get('count', 0)
                narrative += f"**{i}. {theme_name.title()}**\n"
                narrative += f"Tema ini dibahas dalam {theme_count} studi. "

                # Find papers related to this theme
                related_papers = [p for p in papers if theme_name.lower() in
                                 (p.get('title', '') + ' ' + str(p.get('keywords', ''))).lower()][:3]

                if related_papers:
                    authors = [p.get('authors', [''])[0].split(',')[0] for p in related_papers]
                    years = [str(p.get('year', '')) for p in related_papers]
                    citations = [f"{a} ({y})" for a, y in zip(authors, years) if a and y]
                    if citations:
                        narrative += f"Studi-studi yang membahas tema ini antara lain {', '.join(citations[:3])}. "

                narrative += "\n\n"
        else:
            narrative += """Dari analisis yang dilakukan, temuan-temuan studi dapat dikelompokkan berdasarkan beberapa kategori utama yang relevan dengan pertanyaan penelitian. Pembahasan lebih lanjut mengenai masing-masing kategori disajikan dalam subbab berikut.

"""

        # Add TL;DR summaries
        papers_with_tldr = [p for p in papers if p.get('tldr')]
        if papers_with_tldr:
            narrative += "**Ringkasan Temuan Utama**\n\n"
            for i, paper in enumerate(papers_with_tldr[:5], 1):
                author = paper.get('authors', ['Penulis'])[0].split(',')[0] if paper.get('authors') else 'Penulis'
                year = paper.get('year', '')
                tldr = paper.get('tldr', '')
                narrative += f"{i}. {author} ({year}): {tldr}\n\n"

        return narrative

    async def generate_discussion_narrative(
        self,
        slr_results: Dict
    ) -> GeneratedNarrative:
        """Generate discussion section narrative."""

        papers = slr_results.get('acquired_papers', [])
        research_question = slr_results.get('research_question', '')

        if self.anthropic_client:
            narrative = await self._generate_with_llm(
                section="Discussion",
                data={
                    'research_question': research_question,
                    'total_papers': len(papers),
                    'papers': papers[:10],
                    'prisma_stats': slr_results.get('prisma_stats', {})
                },
                template=self._get_discussion_template()
            )
        else:
            narrative = self._generate_discussion_fallback(papers, research_question)

        return GeneratedNarrative(
            section=NarrativeSection.DISCUSSION,
            title="4.5 Diskusi",
            content=narrative,
            word_count=len(narrative.split())
        )

    def _generate_discussion_fallback(
        self,
        papers: List[Dict],
        research_question: str
    ) -> str:
        """Generate discussion without LLM."""

        narrative = f"""Tinjauan sistematis ini bertujuan untuk menjawab pertanyaan penelitian: "{research_question}"

Berdasarkan analisis terhadap {len(papers)} studi yang memenuhi kriteria inklusi, beberapa temuan penting dapat diidentifikasi.

**Temuan Utama**

Hasil sintesis menunjukkan bahwa topik penelitian ini telah mendapat perhatian yang signifikan dalam literatur akademik. Studi-studi yang dianalisis memberikan perspektif yang beragam terhadap permasalahan yang diteliti.

**Implikasi Teoretis**

Temuan tinjauan sistematis ini berkontribusi terhadap pengembangan pemahaman teoretis di bidang terkait. Hasil analisis mengkonfirmasi beberapa konsep yang telah ada sekaligus mengidentifikasi area yang memerlukan pengembangan lebih lanjut.

**Implikasi Praktis**

Dari perspektif praktis, temuan ini memiliki implikasi penting bagi praktisi dan pembuat kebijakan. Rekomendasi yang dapat diberikan berdasarkan temuan meliputi:

1. Perlunya pendekatan yang lebih sistematis dalam implementasi
2. Pentingnya mempertimbangkan konteks lokal dalam penerapan
3. Kebutuhan akan pengembangan kapasitas sumber daya manusia

**Perbandingan dengan Studi Sebelumnya**

Temuan tinjauan sistematis ini sejalan dengan beberapa studi sebelumnya yang juga menemukan pola serupa. Namun demikian, terdapat beberapa perbedaan yang perlu dicermati, terutama terkait dengan konteks dan metodologi yang digunakan.

"""
        return narrative

    async def generate_limitations_narrative(
        self,
        slr_results: Dict
    ) -> GeneratedNarrative:
        """Generate limitations section narrative."""

        prisma_stats = slr_results.get('prisma_stats', {})
        not_retrieved = prisma_stats.get('not_retrieved', 0)

        virtual_fulltext_count = len([
            p for p in slr_results.get('acquired_papers', [])
            if p.get('full_text_source') == 'virtual_fulltext'
        ])

        if self.anthropic_client:
            narrative = await self._generate_with_llm(
                section="Limitations",
                data={
                    'not_retrieved': not_retrieved,
                    'virtual_fulltext_count': virtual_fulltext_count,
                    'total_papers': len(slr_results.get('acquired_papers', []))
                },
                template=self._get_limitations_template()
            )
        else:
            narrative = self._generate_limitations_fallback(
                not_retrieved, virtual_fulltext_count
            )

        return GeneratedNarrative(
            section=NarrativeSection.LIMITATIONS,
            title="4.6 Keterbatasan Studi",
            content=narrative,
            word_count=len(narrative.split())
        )

    def _generate_limitations_fallback(
        self,
        not_retrieved: int,
        virtual_fulltext_count: int
    ) -> str:
        """Generate limitations without LLM."""

        narrative = """Tinjauan sistematis ini memiliki beberapa keterbatasan yang perlu dipertimbangkan dalam menginterpretasikan temuan.

**Keterbatasan Pencarian**

Pertama, pencarian literatur hanya dilakukan pada basis data elektronik tertentu, sehingga kemungkinan terdapat studi relevan yang tidak teridentifikasi (publication bias). Kedua, pembatasan bahasa publikasi ke dalam bahasa Inggris dan Indonesia dapat mengeksklusi studi-studi relevan dalam bahasa lain.

"""

        if not_retrieved > 0:
            narrative += f"""**Keterbatasan Akses**

Sebanyak {not_retrieved} artikel tidak dapat diperoleh teks lengkapnya karena kendala akses, yang dapat mempengaruhi kelengkapan tinjauan ini.

"""

        if virtual_fulltext_count > 0:
            narrative += f"""**Keterbatasan Virtual Full-Text**

Sebanyak {virtual_fulltext_count} artikel dianalisis menggunakan metode Virtual Full-Text (sintesis dari abstrak dan konteks sitasi) karena tidak tersedianya teks lengkap. Hal ini dapat membatasi kedalaman analisis terhadap studi-studi tersebut.

"""

        narrative += """**Keterbatasan Metodologis**

Heterogenitas dalam desain dan metodologi studi yang diinklusi membatasi kemampuan untuk melakukan meta-analisis kuantitatif. Oleh karena itu, sintesis dilakukan secara naratif yang mungkin memiliki tingkat subjektivitas tertentu.

**Keterbatasan Temporal**

Pencarian literatur dilakukan hingga tanggal tertentu, sehingga studi-studi yang dipublikasikan setelah periode pencarian tidak termasuk dalam tinjauan ini.

Meskipun terdapat keterbatasan-keterbatasan tersebut, tinjauan sistematis ini telah dilakukan dengan mengikuti pedoman PRISMA 2020 untuk memastikan transparansi dan reprodusibilitas proses."""

        return narrative

    async def _generate_with_llm(
        self,
        section: str,
        data: Dict,
        template: str
    ) -> str:
        """Generate narrative using LLM."""

        prompt = f"""{template}

DATA:
{self._format_data_for_llm(data)}

INSTRUKSI:
- Tulis dalam bahasa Indonesia yang formal dan akademis
- Gunakan kalimat pasif dan gaya penulisan ilmiah
- Sertakan angka dan statistik yang relevan
- Hindari penggunaan kata ganti orang pertama
- Pastikan koherensi dan alur logis
- Panjang sekitar {self.config.max_section_length} kata

Tulis narasi untuk bagian {section}:"""

        try:
            response = await self.anthropic_client.messages.create(
                model="claude-3-5-sonnet-20241022",
                max_tokens=2000,
                messages=[{"role": "user", "content": prompt}]
            )
            return response.content[0].text
        except Exception as e:
            logger.error(f"LLM generation error: {e}")
            return f"[Error generating narrative: {e}]"

    def _format_data_for_llm(self, data: Dict) -> str:
        """Format data dictionary for LLM prompt."""
        lines = []
        for key, value in data.items():
            if isinstance(value, list):
                if len(value) > 5:
                    value = value[:5]  # Limit list length
            lines.append(f"- {key}: {value}")
        return "\n".join(lines)

    def _get_prisma_template(self) -> str:
        return """Anda adalah penulis akademik yang ahli dalam penulisan tinjauan sistematis.
Tugas: Menulis narasi untuk bagian PRISMA Flow dalam bahasa Indonesia formal.
Format: Paragraf akademis yang menjelaskan proses seleksi studi."""

    def _get_characteristics_template(self) -> str:
        return """Anda adalah penulis akademik yang ahli dalam penulisan tinjauan sistematis.
Tugas: Menulis narasi karakteristik studi yang diinklusi dalam bahasa Indonesia formal.
Format: Paragraf akademis yang merangkum karakteristik studi (tahun, jurnal, metode, negara)."""

    def _get_quality_template(self) -> str:
        return """Anda adalah penulis akademik yang ahli dalam penulisan tinjauan sistematis.
Tugas: Menulis narasi penilaian kualitas studi dalam bahasa Indonesia formal.
Format: Paragraf akademis yang menjelaskan hasil penilaian kualitas metodologis."""

    def _get_thematic_template(self) -> str:
        return """Anda adalah penulis akademik yang ahli dalam penulisan tinjauan sistematis.
Tugas: Menulis narasi sintesis tematik dalam bahasa Indonesia formal.
Format: Paragraf akademis yang menjelaskan tema-tema utama dari temuan studi."""

    def _get_discussion_template(self) -> str:
        return """Anda adalah penulis akademik yang ahli dalam penulisan tinjauan sistematis.
Tugas: Menulis narasi diskusi dalam bahasa Indonesia formal.
Format: Paragraf akademis yang mendiskusikan temuan, implikasi, dan perbandingan dengan studi lain."""

    def _get_limitations_template(self) -> str:
        return """Anda adalah penulis akademik yang ahli dalam penulisan tinjauan sistematis.
Tugas: Menulis narasi keterbatasan studi dalam bahasa Indonesia formal.
Format: Paragraf akademis yang menjelaskan keterbatasan tinjauan sistematis."""

    def export_to_markdown(self) -> str:
        """Export generated narrative to Markdown format."""

        if not self.generated_sections:
            return "# Hasil dan Pembahasan\n\nBelum ada narasi yang dihasilkan."

        md_content = """# BAB IV
# HASIL DAN PEMBAHASAN

"""
        for section in NarrativeSection:
            if section in self.generated_sections:
                narrative = self.generated_sections[section]
                md_content += f"## {narrative.title}\n\n"
                md_content += f"{narrative.content}\n\n"

                # Add tables if any
                for table in narrative.tables:
                    md_content += f"\n{table}\n\n"

        return md_content

    def export_to_word(self, filepath: str) -> bool:
        """Export generated narrative to Word document."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            return False

        doc = Document()

        # Title
        title = doc.add_heading('BAB IV', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_heading('HASIL DAN PEMBAHASAN', level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add sections
        for section in NarrativeSection:
            if section in self.generated_sections:
                narrative = self.generated_sections[section]

                # Section heading
                doc.add_heading(narrative.title, level=2)

                # Content paragraphs
                paragraphs = narrative.content.split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        p = doc.add_paragraph(para.strip())
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        try:
            doc.save(filepath)
            logger.info(f"Word document saved: {filepath}")
            return True
        except Exception as e:
            logger.error(f"Error saving Word document: {e}")
            return False


# Convenience function
async def generate_results_chapter(
    slr_results: Dict,
    anthropic_api_key: str = None,
    output_format: str = "markdown"
) -> str:
    """
    Generate Results and Discussion chapter from SLR results.

    Args:
        slr_results: Dictionary containing SLR results
        anthropic_api_key: Optional API key for LLM enhancement
        output_format: "markdown" or "word"

    Returns:
        Generated narrative as string (markdown) or filepath (word)
    """
    anthropic_client = None
    if anthropic_api_key:
        try:
            from anthropic import AsyncAnthropic
            anthropic_client = AsyncAnthropic(api_key=anthropic_api_key)
        except ImportError:
            logger.warning("Anthropic client not available")

    generator = NarrativeGenerator(anthropic_client=anthropic_client)
    await generator.generate_full_chapter(slr_results)

    if output_format == "markdown":
        return generator.export_to_markdown()
    elif output_format == "word":
        filepath = f"hasil_pembahasan_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        generator.export_to_word(filepath)
        return filepath

    return generator.export_to_markdown()
