"""
BiblioAgent AI - DOCX Report Generator
=======================================
Generates professional Word documents for Systematic Literature Review reports.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from typing import Dict, List, Optional
import re
from datetime import datetime


class DocxGenerator:
    """
    Generate professional SLR reports in DOCX format.

    Features:
    - 5-chapter structure (BAB I - BAB V)
    - Auto-formatted bibliography
    - Professional styling
    - Table of contents ready
    """

    def __init__(self, researcher_name: str = "Peneliti", institution: str = ""):
        """
        Initialize DocxGenerator.

        Args:
            researcher_name: Name of the researcher/author
            institution: Institution name (optional)
        """
        self.doc = Document()
        self.researcher_name = researcher_name
        self.institution = institution
        self._setup_styles()

    def _setup_styles(self):
        """Setup document styles for professional appearance."""
        # Set default font
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        # Paragraph formatting
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = 1.5
        paragraph_format.space_after = Pt(6)

    def add_heading(self, text: str, level: int):
        """
        Add a heading with proper formatting.

        Args:
            text: Heading text
            level: Heading level (0=Title, 1=Chapter, 2=Section)
        """
        heading = self.doc.add_heading(text, level=level)

        # Center align for title and chapter headings
        if level <= 1:
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Set font for heading
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run.font.bold = True
            if level == 0:
                run.font.size = Pt(16)
            elif level == 1:
                run.font.size = Pt(14)
            else:
                run.font.size = Pt(12)

    def add_paragraph(self, text: str, justify: bool = True, first_line_indent: bool = True):
        """
        Add a paragraph with proper formatting.

        Args:
            text: Paragraph text
            justify: Whether to justify text alignment
            first_line_indent: Whether to add first line indent
        """
        p = self.doc.add_paragraph(text)

        if justify:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if first_line_indent:
            p.paragraph_format.first_line_indent = Cm(1.27)  # 0.5 inch

        # Set font
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        return p

    def add_title_page(self, title: str, subtitle: str = ""):
        """
        Add a professional title page.

        Args:
            title: Main report title
            subtitle: Optional subtitle
        """
        # Add some spacing at top
        for _ in range(3):
            self.doc.add_paragraph()

        # Main title
        self.add_heading(title.upper(), 0)

        if subtitle:
            p = self.doc.add_paragraph(subtitle)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)

        # Spacing
        for _ in range(5):
            self.doc.add_paragraph()

        # Author info
        p = self.doc.add_paragraph(f"Disusun oleh:")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = self.doc.add_paragraph(self.researcher_name)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(14)

        if self.institution:
            p = self.doc.add_paragraph(self.institution)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Spacing
        for _ in range(5):
            self.doc.add_paragraph()

        # Date
        current_date = datetime.now().strftime("%B %Y")
        p = self.doc.add_paragraph(current_date)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Page break after title
        self.doc.add_page_break()

    def add_chapter_content(self, chapter_title: str, content: str):
        """
        Add a chapter with its content.

        Args:
            chapter_title: Chapter title (e.g., "BAB I PENDAHULUAN")
            content: Chapter content text
        """
        # Add chapter heading
        self.add_heading(chapter_title, 1)

        # Process content - split by sections
        lines = content.split('\n')
        current_section = None
        current_paragraph = []

        for line in lines:
            line = line.strip()

            if not line:
                # Empty line - flush current paragraph
                if current_paragraph:
                    para_text = ' '.join(current_paragraph)
                    self.add_paragraph(para_text)
                    current_paragraph = []
                continue

            # Check for subheading patterns
            # Pattern: "1.1", "1.1.", "A.", "1.1.1", etc.
            subheading_patterns = [
                r'^(\d+\.\d+\.?\d*\.?\s+.+)$',  # 1.1 Title or 1.1.1 Title
                r'^([A-Z]\.\s+.+)$',  # A. Title
                r'^(#+\s+.+)$',  # Markdown style
            ]

            is_subheading = False
            for pattern in subheading_patterns:
                match = re.match(pattern, line)
                if match:
                    # Flush current paragraph
                    if current_paragraph:
                        para_text = ' '.join(current_paragraph)
                        self.add_paragraph(para_text)
                        current_paragraph = []

                    # Add as subheading
                    subheading_text = line.lstrip('#').strip()
                    self.add_heading(subheading_text, 2)
                    is_subheading = True
                    break

            if not is_subheading:
                current_paragraph.append(line)

        # Flush remaining paragraph
        if current_paragraph:
            para_text = ' '.join(current_paragraph)
            self.add_paragraph(para_text)

    def add_bibliography(self, references: List[str], title: str = "DAFTAR PUSTAKA"):
        """
        Add bibliography section.

        Args:
            references: List of reference strings
            title: Section title
        """
        self.doc.add_page_break()
        self.add_heading(title, 1)

        # Sort references alphabetically
        sorted_refs = sorted(references, key=lambda x: x.lower())

        for ref in sorted_refs:
            p = self.doc.add_paragraph(ref, style='List Bullet')
            p.paragraph_format.left_indent = Cm(1.27)
            p.paragraph_format.first_line_indent = Cm(-0.63)  # Hanging indent

            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

    def add_table(self, headers: List[str], rows: List[List[str]], title: str = ""):
        """
        Add a table to the document.

        Args:
            headers: List of column headers
            rows: List of row data (each row is a list of cell values)
            title: Optional table title
        """
        if title:
            p = self.doc.add_paragraph(title)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True

        # Create table
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        # Add headers
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)

        # Add data rows
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, cell_value in enumerate(row_data):
                row_cells[i].text = str(cell_value)
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)

        # Add spacing after table
        self.doc.add_paragraph()

    def generate_report(
        self,
        chapters: Dict[str, str],
        bibliography: List[str],
        filename: str = "Laporan_SLR_BiblioAgent.docx",
        title: str = "LAPORAN SYSTEMATIC LITERATURE REVIEW",
        subtitle: str = "",
        include_title_page: bool = True
    ) -> str:
        """
        Generate complete SLR report.

        Args:
            chapters: Dictionary of chapter titles and contents
            bibliography: List of reference strings
            filename: Output filename
            title: Report title
            subtitle: Report subtitle
            include_title_page: Whether to include title page

        Returns:
            Path to generated file
        """
        # Reset document
        self.doc = Document()
        self._setup_styles()

        # Title page
        if include_title_page:
            self.add_title_page(title, subtitle)
        else:
            self.add_heading(title, 0)
            p = self.doc.add_paragraph(f"Disusun oleh: {self.researcher_name}")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.doc.add_paragraph()

        # Chapter iteration (BAB I - BAB V)
        chapter_order = [
            "BAB_I_PENDAHULUAN",
            "BAB_II_TINJAUAN_PUSTAKA",
            "BAB_III_METODOLOGI",
            "BAB_IV_HASIL_PEMBAHASAN",
            "BAB_V_KESIMPULAN"
        ]

        # Try ordered first, then fall back to dict order
        processed_chapters = set()

        for chapter_key in chapter_order:
            # Try exact match
            if chapter_key in chapters:
                chapter_title = chapter_key.replace("_", " ")
                self.add_chapter_content(chapter_title, chapters[chapter_key])
                processed_chapters.add(chapter_key)
                self.doc.add_page_break()
            else:
                # Try fuzzy match
                for key, content in chapters.items():
                    normalized_key = key.upper().replace(" ", "_")
                    if chapter_key in normalized_key or normalized_key in chapter_key:
                        chapter_title = key.replace("_", " ")
                        self.add_chapter_content(chapter_title, content)
                        processed_chapters.add(key)
                        self.doc.add_page_break()
                        break

        # Add any remaining chapters not in standard order
        for key, content in chapters.items():
            if key not in processed_chapters:
                chapter_title = key.replace("_", " ")
                self.add_chapter_content(chapter_title, content)
                self.doc.add_page_break()

        # Bibliography
        if bibliography:
            self.add_bibliography(bibliography)

        # Save document
        self.doc.save(filename)
        print(f"Success! Dokumen {filename} telah siap.")

        return filename

    def generate_from_orchestrator(
        self,
        orchestrator_result: Dict,
        bibliography: List[str],
        filename: str = "Laporan_SLR_BiblioAgent.docx",
        title: str = "LAPORAN SYSTEMATIC LITERATURE REVIEW",
        subtitle: str = ""
    ) -> str:
        """
        Generate report from NarrativeOrchestrator result.

        Args:
            orchestrator_result: Result from NarrativeOrchestrator.generate_full_report()
            bibliography: List of reference strings
            filename: Output filename
            title: Report title
            subtitle: Report subtitle

        Returns:
            Path to generated file
        """
        chapters = {}

        # Map orchestrator keys to chapter format
        key_mapping = {
            'bab_1': 'BAB_I_PENDAHULUAN',
            'bab_2': 'BAB_II_TINJAUAN_PUSTAKA',
            'bab_3': 'BAB_III_METODOLOGI',
            'bab_4': 'BAB_IV_HASIL_PEMBAHASAN',
            'bab_5': 'BAB_V_KESIMPULAN'
        }

        for old_key, new_key in key_mapping.items():
            if old_key in orchestrator_result:
                chapters[new_key] = orchestrator_result[old_key]

        return self.generate_report(
            chapters=chapters,
            bibliography=bibliography,
            filename=filename,
            title=title,
            subtitle=subtitle
        )


def generate_slr_docx(
    chapters: Dict[str, str],
    bibliography: List[str],
    researcher_name: str = "Peneliti",
    institution: str = "",
    filename: str = "Laporan_SLR_BiblioAgent.docx",
    title: str = "LAPORAN SYSTEMATIC LITERATURE REVIEW"
) -> str:
    """
    Convenience function to generate SLR report.

    Args:
        chapters: Dictionary of chapter titles and contents
        bibliography: List of reference strings
        researcher_name: Author name
        institution: Institution name
        filename: Output filename
        title: Report title

    Returns:
        Path to generated file
    """
    generator = DocxGenerator(
        researcher_name=researcher_name,
        institution=institution
    )
    return generator.generate_report(
        chapters=chapters,
        bibliography=bibliography,
        filename=filename,
        title=title
    )
