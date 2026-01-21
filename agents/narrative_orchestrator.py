"""
Narrative Orchestrator - Full Research Report Generator
========================================================
Generates complete 5-chapter research report in formal Indonesian academic style.

Chapters:
- Bab 1: Pendahuluan (Background, urgency, research gap)
- Bab 2: Tinjauan Pustaka (Literature synthesis by thematic clusters)
- Bab 3: Metodologi (SLR methodology, AI screening, Waterfall Retrieval)
- Bab 4: Hasil dan Pembahasan (Data extraction analysis)
- Bab 5: Kesimpulan dan Saran (Conclusions and recommendations)
"""

import os
import logging
from typing import Dict, List, Any, Optional
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum

logger = logging.getLogger(__name__)


class ChapterType(str, Enum):
    """Research report chapter types."""
    BAB_1_PENDAHULUAN = "bab_1"
    BAB_2_TINJAUAN_PUSTAKA = "bab_2"
    BAB_3_METODOLOGI = "bab_3"
    BAB_4_HASIL_PEMBAHASAN = "bab_4"
    BAB_5_KESIMPULAN = "bab_5"


@dataclass
class ChapterContent:
    """Generated chapter content."""
    chapter_type: ChapterType
    title: str
    content: str
    word_count: int = 0
    generated_at: str = field(default_factory=lambda: datetime.now().isoformat())

    def __post_init__(self):
        if self.content:
            self.word_count = len(self.content.split())


@dataclass
class ResearchReportConfig:
    """Configuration for report generation."""
    language: str = "id"  # Indonesian
    style: str = "formal_academic"
    citation_style: str = "APA7"
    include_tables: bool = True
    include_figures: bool = True
    target_word_count: Dict[str, int] = field(default_factory=lambda: {
        "bab_1": 1500,
        "bab_2": 3000,
        "bab_3": 2000,
        "bab_4": 4000,
        "bab_5": 1000,
    })


CHAPTER_TITLES = {
    ChapterType.BAB_1_PENDAHULUAN: "BAB I PENDAHULUAN",
    ChapterType.BAB_2_TINJAUAN_PUSTAKA: "BAB II TINJAUAN PUSTAKA",
    ChapterType.BAB_3_METODOLOGI: "BAB III METODOLOGI PENELITIAN",
    ChapterType.BAB_4_HASIL_PEMBAHASAN: "BAB IV HASIL DAN PEMBAHASAN",
    ChapterType.BAB_5_KESIMPULAN: "BAB V KESIMPULAN DAN SARAN",
}


SYSTEM_PROMPT = """Kamu adalah Senior Research Consultant dengan keahlian dalam penulisan akademik.

PANDUAN PENULISAN:
1. Gunakan bahasa Indonesia formal akademik standar jurnal Q1
2. Struktur paragraf yang logis dan koheren
3. Setiap klaim harus didukung data atau referensi
4. Gunakan terminologi ilmiah yang tepat
5. Hindari pengulangan kata dan kalimat yang redundan
6. Transisi antar paragraf harus smooth dan logis

FORMAT OUTPUT:
- Gunakan heading dan subheading yang jelas
- Paragraf minimal 4-5 kalimat
- Sertakan penomoran untuk list items
- Gunakan format markdown untuk struktur"""


class NarrativeOrchestrator:
    """
    Orchestrates generation of complete research report chapters.

    Uses Claude LLM to generate formal Indonesian academic text
    based on SLR data and extraction tables.
    """

    def __init__(
        self,
        api_key: str = None,
        config: ResearchReportConfig = None,
        use_langchain: bool = True
    ):
        """
        Initialize the orchestrator.

        Args:
            api_key: Anthropic API key (defaults to env var)
            config: Report generation configuration
            use_langchain: Whether to use LangChain (True) or direct Anthropic client (False)
        """
        self.api_key = api_key or os.getenv("ANTHROPIC_API_KEY")
        self.config = config or ResearchReportConfig()
        self.use_langchain = use_langchain
        self.llm = None
        self.chapters: Dict[ChapterType, ChapterContent] = {}

        self._initialize_llm()

    def _initialize_llm(self):
        """Initialize the LLM client."""
        if not self.api_key:
            logger.warning("No API key provided. LLM features will be disabled.")
            return

        if self.use_langchain:
            try:
                from langchain_anthropic import ChatAnthropic
                self.llm = ChatAnthropic(
                    model="claude-sonnet-4-20250514",
                    api_key=self.api_key,
                    max_tokens=4096,
                    temperature=0.3
                )
                logger.info("LangChain ChatAnthropic initialized")
            except ImportError:
                logger.warning("langchain_anthropic not installed, falling back to direct client")
                self.use_langchain = False

        if not self.use_langchain:
            try:
                from anthropic import Anthropic
                self.llm = Anthropic(api_key=self.api_key)
                logger.info("Direct Anthropic client initialized")
            except ImportError:
                logger.error("Neither langchain_anthropic nor anthropic package available")
                self.llm = None

    def _invoke_llm(self, instruction: str, context: str) -> str:
        """
        Invoke LLM with instruction and context.

        Args:
            instruction: Specific instruction for the chapter
            context: Supporting data and context

        Returns:
            Generated text content
        """
        if not self.llm:
            return self._generate_template_content(instruction)

        full_prompt = f"Instruksi: {instruction}\n\nData Pendukung:\n{context}"

        try:
            if self.use_langchain:
                from langchain.schema import SystemMessage, HumanMessage
                response = self.llm.invoke([
                    SystemMessage(content=SYSTEM_PROMPT),
                    HumanMessage(content=full_prompt)
                ])
                return response.content
            else:
                response = self.llm.messages.create(
                    model="claude-sonnet-4-20250514",
                    max_tokens=4096,
                    system=SYSTEM_PROMPT,
                    messages=[{"role": "user", "content": full_prompt}]
                )
                return response.content[0].text

        except Exception as e:
            logger.error(f"LLM invocation failed: {e}")
            return self._generate_template_content(instruction)

    def _generate_template_content(self, instruction: str) -> str:
        """Generate template content when LLM is unavailable."""
        return f"""[Konten akan digenerate berdasarkan instruksi berikut]

{instruction}

---
*Catatan: Konten ini adalah template. Untuk hasil optimal, pastikan API key telah dikonfigurasi.*
"""

    def generate_bab_1_pendahuluan(
        self,
        research_question: str,
        scopus_metadata: Dict[str, Any],
        background_context: str = ""
    ) -> ChapterContent:
        """
        Generate BAB I - Pendahuluan (Introduction).

        Sections:
        - 1.1 Latar Belakang Masalah
        - 1.2 Rumusan Masalah
        - 1.3 Tujuan Penelitian
        - 1.4 Manfaat Penelitian
        - 1.5 Batasan Penelitian
        """
        # Prepare Scopus statistics
        total_papers = scopus_metadata.get('total_results', 0)
        year_range = scopus_metadata.get('year_range', 'tidak diketahui')
        top_sources = scopus_metadata.get('top_sources', [])
        trend_data = scopus_metadata.get('publication_trend', {})

        stats_summary = f"""
Statistik Pencarian Scopus:
- Total publikasi teridentifikasi: {total_papers}
- Rentang tahun: {year_range}
- Sumber publikasi utama: {', '.join(top_sources[:5]) if top_sources else 'N/A'}
- Tren publikasi: {self._format_trend(trend_data)}
"""

        instruction = f"""Susun BAB I PENDAHULUAN dengan struktur berikut:

1.1 Latar Belakang Masalah
- Jelaskan urgensi topik penelitian berdasarkan data statistik Scopus
- Identifikasi research gap yang ada
- Gunakan data tren publikasi untuk menunjukkan perkembangan bidang ini
- Cantumkan statistik global yang relevan

1.2 Rumusan Masalah
- Formulasikan pertanyaan penelitian utama: "{research_question}"
- Turunkan sub-pertanyaan penelitian yang spesifik

1.3 Tujuan Penelitian
- Tujuan umum yang selaras dengan rumusan masalah
- Tujuan khusus yang terukur dan spesifik

1.4 Manfaat Penelitian
- Manfaat teoretis bagi pengembangan ilmu
- Manfaat praktis bagi stakeholder terkait

1.5 Batasan Penelitian
- Batasan lingkup, waktu, dan metodologi

Konteks tambahan: {background_context}"""

        content = self._invoke_llm(instruction, stats_summary)

        chapter = ChapterContent(
            chapter_type=ChapterType.BAB_1_PENDAHULUAN,
            title=CHAPTER_TITLES[ChapterType.BAB_1_PENDAHULUAN],
            content=content
        )
        self.chapters[ChapterType.BAB_1_PENDAHULUAN] = chapter
        return chapter

    def generate_bab_2_tinjauan_pustaka(
        self,
        research_question: str,
        papers: List[Dict[str, Any]],
        thematic_clusters: Dict[str, List[Dict]] = None
    ) -> ChapterContent:
        """
        Generate BAB II - Tinjauan Pustaka (Literature Review).

        Sections:
        - 2.1 Landasan Teori
        - 2.2 Kajian Penelitian Terdahulu
        - 2.3 Sintesis Literatur
        - 2.4 Kerangka Konseptual
        """
        # Organize papers by clusters if available
        if thematic_clusters:
            cluster_summary = self._format_thematic_clusters(thematic_clusters)
        else:
            cluster_summary = self._auto_cluster_papers(papers)

        # Extract key theories and concepts
        papers_summary = self._summarize_papers_for_literature(papers[:20])  # Top 20 papers

        instruction = f"""Susun BAB II TINJAUAN PUSTAKA dengan struktur:

2.1 Landasan Teori
- Definisi dan konsep kunci terkait topik penelitian
- Teori-teori utama yang mendasari penelitian
- Evolusi pemikiran dalam bidang ini

2.2 Kajian Penelitian Terdahulu
- Sintesis hasil penelitian berdasarkan klaster tematik
- Bandingkan perspektif dan temuan antar kelompok penelitian
- Identifikasi konsensus dan kontroversi dalam literatur

2.3 Sintesis Literatur
- Integrasikan temuan dari berbagai studi
- Identifikasi pola dan tren dalam literatur
- Highlight gaps yang belum terisi

2.4 Kerangka Konseptual
- Bangun kerangka berdasarkan sintesis literatur
- Tunjukkan hubungan antar variabel/konsep
- Jelaskan bagaimana kerangka ini menjawab research question

Pertanyaan Penelitian: {research_question}

Klaster Tematik:
{cluster_summary}"""

        content = self._invoke_llm(instruction, papers_summary)

        chapter = ChapterContent(
            chapter_type=ChapterType.BAB_2_TINJAUAN_PUSTAKA,
            title=CHAPTER_TITLES[ChapterType.BAB_2_TINJAUAN_PUSTAKA],
            content=content
        )
        self.chapters[ChapterType.BAB_2_TINJAUAN_PUSTAKA] = chapter
        return chapter

    def generate_bab_3_metodologi(
        self,
        prisma_stats: Dict[str, int],
        search_strategy: Dict[str, Any] = None,
        screening_details: Dict[str, Any] = None
    ) -> ChapterContent:
        """
        Generate BAB III - Metodologi Penelitian.

        Sections:
        - 3.1 Desain Penelitian
        - 3.2 Strategi Pencarian Literatur
        - 3.3 Kriteria Seleksi
        - 3.4 Proses Screening
        - 3.5 Ekstraksi dan Analisis Data
        - 3.6 Penilaian Kualitas
        """
        # Format PRISMA statistics
        prisma_summary = f"""
Statistik PRISMA:
- Artikel teridentifikasi: {prisma_stats.get('identified', 0)}
- Duplikat dihapus: {prisma_stats.get('duplicates_removed', 0)}
- Artikel di-screening: {prisma_stats.get('screened', 0)}
- Dieksklusi saat screening: {prisma_stats.get('excluded_screening', 0)}
- Artikel untuk retrieval: {prisma_stats.get('sought_retrieval', 0)}
- Tidak dapat diakses: {prisma_stats.get('not_retrieved', 0)}
- Dinilai eligibilitas: {prisma_stats.get('assessed_eligibility', 0)}
- Dieksklusi saat eligibilitas: {prisma_stats.get('excluded_eligibility', 0)}
- Diinklusi dalam sintesis: {prisma_stats.get('included_synthesis', 0)}
"""

        search_info = ""
        if search_strategy:
            search_info = f"""
Strategi Pencarian:
- Database: {', '.join(search_strategy.get('databases', ['Scopus']))}
- Query: {search_strategy.get('query', 'N/A')}
- Rentang tahun: {search_strategy.get('date_range', 'N/A')}
"""

        instruction = f"""Susun BAB III METODOLOGI PENELITIAN dengan struktur:

3.1 Desain Penelitian
- Jelaskan pendekatan Systematic Literature Review (SLR)
- Rujuk pedoman PRISMA 2020
- Jelaskan rasionale pemilihan metode ini

3.2 Strategi Pencarian Literatur
- Database yang digunakan dan justifikasinya
- Kata kunci dan Boolean operators
- Pembatasan bahasa dan periode waktu

3.3 Kriteria Seleksi
- Kriteria inklusi (dengan justifikasi)
- Kriteria eksklusi (dengan justifikasi)
- Definisi operasional kriteria

3.4 Proses Screening dengan AI
- Jelaskan penggunaan AI (Claude) untuk screening judul/abstrak
- Proses 4-fase screening dengan confidence scoring
- Validasi hasil screening AI

3.5 Strategi Waterfall Retrieval
- Jelaskan cascading retrieval: Semantic Scholar -> Unpaywall -> CORE -> ArXiv
- Virtual Full-Text synthesis untuk paper yang tidak dapat diakses
- Proses validasi dan quality checking

3.6 Ekstraksi dan Analisis Data
- Template ekstraksi data
- Metode sintesis (naratif/tematik)
- Penilaian kualitas dengan JBI Critical Appraisal

{search_info}"""

        content = self._invoke_llm(instruction, prisma_summary)

        chapter = ChapterContent(
            chapter_type=ChapterType.BAB_3_METODOLOGI,
            title=CHAPTER_TITLES[ChapterType.BAB_3_METODOLOGI],
            content=content
        )
        self.chapters[ChapterType.BAB_3_METODOLOGI] = chapter
        return chapter

    def generate_bab_4_hasil_pembahasan(
        self,
        research_question: str,
        extraction_table: List[Dict[str, Any]],
        quality_scores: List[Dict[str, Any]] = None,
        themes: List[str] = None
    ) -> ChapterContent:
        """
        Generate BAB IV - Hasil dan Pembahasan.

        Sections:
        - 4.1 Proses Seleksi Studi (PRISMA)
        - 4.2 Karakteristik Studi
        - 4.3 Penilaian Kualitas
        - 4.4 Sintesis Temuan
        - 4.5 Pembahasan
        """
        # Format extraction table
        table_summary = self._format_extraction_table(extraction_table)

        # Quality summary
        quality_summary = ""
        if quality_scores:
            high = sum(1 for q in quality_scores if q.get('category') == 'HIGH')
            moderate = sum(1 for q in quality_scores if q.get('category') == 'MODERATE')
            low = sum(1 for q in quality_scores if q.get('category') == 'LOW')
            quality_summary = f"""
Distribusi Kualitas:
- Kualitas Tinggi: {high} studi
- Kualitas Sedang: {moderate} studi
- Kualitas Rendah: {low} studi
"""

        themes_text = ""
        if themes:
            themes_text = f"Tema yang teridentifikasi: {', '.join(themes)}"

        instruction = f"""Susun BAB IV HASIL DAN PEMBAHASAN dengan struktur:

4.1 Proses Seleksi Studi (PRISMA Flow)
- Narasikan proses seleksi sesuai diagram PRISMA
- Jelaskan alasan eksklusi pada setiap tahap
- Sertakan statistik pada setiap fase

4.2 Karakteristik Studi yang Diinklusi
- Deskripsi demografis studi (tahun, negara, jurnal)
- Desain penelitian yang digunakan
- Populasi dan sampel
- Tabel ringkasan karakteristik

4.3 Penilaian Kualitas Metodologis
- Hasil penilaian dengan JBI tools
- Distribusi skor kualitas
- Implikasi terhadap sintesis

4.4 Sintesis Temuan Utama
- Organisasi temuan berdasarkan tema
- Hubungkan dengan pertanyaan penelitian
- Sertakan data kuantitatif jika relevan
{themes_text}

4.5 Pembahasan
- Interpretasi temuan dalam konteks literatur
- Bandingkan dengan penelitian sebelumnya
- Implikasi teoretis dan praktis
- Kekuatan dan keterbatasan studi

Pertanyaan Penelitian: {research_question}
{quality_summary}"""

        content = self._invoke_llm(instruction, table_summary)

        chapter = ChapterContent(
            chapter_type=ChapterType.BAB_4_HASIL_PEMBAHASAN,
            title=CHAPTER_TITLES[ChapterType.BAB_4_HASIL_PEMBAHASAN],
            content=content
        )
        self.chapters[ChapterType.BAB_4_HASIL_PEMBAHASAN] = chapter
        return chapter

    def generate_bab_5_kesimpulan(
        self,
        research_question: str,
        key_findings: List[str] = None,
        practical_implications: List[str] = None
    ) -> ChapterContent:
        """
        Generate BAB V - Kesimpulan dan Saran.

        Sections:
        - 5.1 Kesimpulan
        - 5.2 Implikasi
        - 5.3 Saran dan Rekomendasi
        - 5.4 Agenda Penelitian Mendatang
        """
        findings_text = ""
        if key_findings:
            findings_text = "Temuan Utama:\n" + "\n".join(f"- {f}" for f in key_findings)

        implications_text = ""
        if practical_implications:
            implications_text = "Implikasi Praktis:\n" + "\n".join(f"- {i}" for i in practical_implications)

        # Get summary from previous chapters if available
        bab4_summary = ""
        if ChapterType.BAB_4_HASIL_PEMBAHASAN in self.chapters:
            bab4_content = self.chapters[ChapterType.BAB_4_HASIL_PEMBAHASAN].content
            bab4_summary = f"Ringkasan Bab IV:\n{bab4_content[:2000]}..."

        instruction = f"""Susun BAB V KESIMPULAN DAN SARAN dengan struktur:

5.1 Kesimpulan
- Jawab pertanyaan penelitian secara langsung dan ringkas
- Rangkum temuan utama (3-5 poin kunci)
- Sintesis kontribusi penelitian

5.2 Implikasi Penelitian
- Implikasi teoretis: kontribusi terhadap body of knowledge
- Implikasi praktis: rekomendasi untuk praktisi/pembuat kebijakan
- Implikasi metodologis: kontribusi terhadap metode penelitian

5.3 Saran dan Rekomendasi
- Rekomendasi berbasis bukti untuk stakeholder
- Rekomendasi kebijakan yang actionable
- Prioritaskan berdasarkan urgensi dan feasibility

5.4 Agenda Penelitian Mendatang
- Identifikasi research gaps yang masih terbuka
- Usulan topik penelitian lanjutan
- Saran metodologi untuk penelitian future

Pertanyaan Penelitian: {research_question}
{findings_text}
{implications_text}"""

        content = self._invoke_llm(instruction, bab4_summary)

        chapter = ChapterContent(
            chapter_type=ChapterType.BAB_5_KESIMPULAN,
            title=CHAPTER_TITLES[ChapterType.BAB_5_KESIMPULAN],
            content=content
        )
        self.chapters[ChapterType.BAB_5_KESIMPULAN] = chapter
        return chapter

    def generate_full_report(
        self,
        research_question: str,
        scopus_metadata: Dict[str, Any],
        extraction_table: List[Dict[str, Any]],
        papers: List[Dict[str, Any]] = None,
        prisma_stats: Dict[str, int] = None,
        thematic_clusters: Dict[str, List[Dict]] = None,
        quality_scores: List[Dict[str, Any]] = None
    ) -> Dict[str, ChapterContent]:
        """
        Generate complete research report with all 5 chapters.

        Args:
            research_question: Main research question
            scopus_metadata: Scopus search statistics and metadata
            extraction_table: Data extraction results
            papers: List of included papers
            prisma_stats: PRISMA flow statistics
            thematic_clusters: Papers organized by themes
            quality_scores: Quality assessment results

        Returns:
            Dictionary of generated chapters
        """
        papers = papers or []
        prisma_stats = prisma_stats or {}

        logger.info("Starting full report generation...")

        # Generate each chapter sequentially
        print("Generating Bab 1 - Pendahuluan...")
        self.generate_bab_1_pendahuluan(research_question, scopus_metadata)

        print("Generating Bab 2 - Tinjauan Pustaka...")
        self.generate_bab_2_tinjauan_pustaka(research_question, papers, thematic_clusters)

        print("Generating Bab 3 - Metodologi...")
        self.generate_bab_3_metodologi(prisma_stats)

        print("Generating Bab 4 - Hasil dan Pembahasan...")
        self.generate_bab_4_hasil_pembahasan(
            research_question,
            extraction_table,
            quality_scores
        )

        print("Generating Bab 5 - Kesimpulan...")
        # Extract key findings from Bab 4
        key_findings = self._extract_key_findings(extraction_table)
        self.generate_bab_5_kesimpulan(research_question, key_findings)

        logger.info(f"Report generation complete. {len(self.chapters)} chapters generated.")
        return self.chapters

    # Helper methods
    def _format_trend(self, trend_data: Dict) -> str:
        """Format publication trend data."""
        if not trend_data:
            return "Data tren tidak tersedia"

        sorted_years = sorted(trend_data.items())
        if len(sorted_years) >= 2:
            first_year, first_count = sorted_years[0]
            last_year, last_count = sorted_years[-1]
            growth = ((last_count - first_count) / first_count * 100) if first_count > 0 else 0
            return f"{first_year}: {first_count} -> {last_year}: {last_count} ({growth:+.1f}%)"
        return str(trend_data)

    def _format_thematic_clusters(self, clusters: Dict[str, List[Dict]]) -> str:
        """Format thematic clusters for prompt."""
        result = []
        for theme, papers in clusters.items():
            paper_titles = [p.get('title', 'Untitled')[:50] for p in papers[:5]]
            result.append(f"**{theme}** ({len(papers)} studi):\n  - " + "\n  - ".join(paper_titles))
        return "\n\n".join(result)

    def _auto_cluster_papers(self, papers: List[Dict]) -> str:
        """Auto-cluster papers by year or simple heuristics."""
        if not papers:
            return "Tidak ada paper untuk di-cluster"

        by_year = {}
        for p in papers:
            year = p.get('year', 'Unknown')
            if year not in by_year:
                by_year[year] = []
            by_year[year].append(p.get('title', 'Untitled'))

        result = []
        for year in sorted(by_year.keys(), reverse=True):
            titles = by_year[year][:3]
            result.append(f"**{year}** ({len(by_year[year])} studi): {', '.join(titles[:2])}...")

        return "\n".join(result)

    def _summarize_papers_for_literature(self, papers: List[Dict]) -> str:
        """Summarize papers for literature review."""
        if not papers:
            return "Tidak ada paper untuk disintesis"

        summaries = []
        for p in papers:
            title = p.get('title', 'Untitled')
            authors = p.get('authors', ['Unknown'])
            year = p.get('year', 'N/A')
            findings = p.get('findings', p.get('abstract', ''))[:200]

            author_str = authors[0] if isinstance(authors, list) else authors
            summaries.append(f"- {author_str} ({year}): {title[:60]}... - {findings}...")

        return "\n".join(summaries[:15])

    def _format_extraction_table(self, table: List[Dict]) -> str:
        """Format extraction table for prompt."""
        if not table:
            return "Tidak ada data ekstraksi"

        formatted = []
        for i, row in enumerate(table[:20], 1):
            entry = f"""
Studi {i}:
- Judul: {row.get('title', 'N/A')[:80]}
- Penulis: {row.get('authors', 'N/A')}
- Tahun: {row.get('year', 'N/A')}
- Desain: {row.get('study_design', 'N/A')}
- Sampel: {row.get('sample_size', 'N/A')}
- Temuan: {row.get('findings', row.get('key_findings', 'N/A'))[:150]}
- Kualitas: {row.get('quality_category', row.get('quality_score', 'N/A'))}
"""
            formatted.append(entry)

        return "\n".join(formatted)

    def _extract_key_findings(self, extraction_table: List[Dict]) -> List[str]:
        """Extract key findings from extraction table."""
        findings = []
        for row in extraction_table[:10]:
            finding = row.get('findings', row.get('key_findings', ''))
            if finding:
                findings.append(finding[:100])
        return findings

    def export_to_markdown(self) -> str:
        """Export all chapters to markdown format."""
        lines = [
            "# LAPORAN PENELITIAN",
            f"*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}*",
            "",
            "---",
            ""
        ]

        chapter_order = [
            ChapterType.BAB_1_PENDAHULUAN,
            ChapterType.BAB_2_TINJAUAN_PUSTAKA,
            ChapterType.BAB_3_METODOLOGI,
            ChapterType.BAB_4_HASIL_PEMBAHASAN,
            ChapterType.BAB_5_KESIMPULAN,
        ]

        for chapter_type in chapter_order:
            if chapter_type in self.chapters:
                chapter = self.chapters[chapter_type]
                lines.extend([
                    f"## {chapter.title}",
                    "",
                    chapter.content,
                    "",
                    f"*Word count: {chapter.word_count}*",
                    "",
                    "---",
                    ""
                ])

        return "\n".join(lines)

    def export_to_word(self, filepath: str) -> bool:
        """Export all chapters to Word document."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.warning("python-docx not installed")
            return False

        doc = Document()

        # Title
        title = doc.add_heading("LAPORAN PENELITIAN", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_page_break()

        chapter_order = [
            ChapterType.BAB_1_PENDAHULUAN,
            ChapterType.BAB_2_TINJAUAN_PUSTAKA,
            ChapterType.BAB_3_METODOLOGI,
            ChapterType.BAB_4_HASIL_PEMBAHASAN,
            ChapterType.BAB_5_KESIMPULAN,
        ]

        for chapter_type in chapter_order:
            if chapter_type in self.chapters:
                chapter = self.chapters[chapter_type]

                # Chapter heading
                doc.add_heading(chapter.title, 1)

                # Content paragraphs
                for para in chapter.content.split('\n\n'):
                    if para.strip():
                        if para.startswith('#'):
                            # Subheading
                            level = para.count('#')
                            text = para.lstrip('#').strip()
                            doc.add_heading(text, min(level + 1, 4))
                        elif para.startswith('**') and para.endswith('**'):
                            # Bold paragraph
                            p = doc.add_paragraph()
                            p.add_run(para.strip('*')).bold = True
                        else:
                            doc.add_paragraph(para.strip())

                doc.add_page_break()

        doc.save(filepath)
        logger.info(f"Word document saved: {filepath}")
        return True


# Convenience function for quick usage
def generate_full_research_report(
    research_question: str,
    scopus_metadata: Dict[str, Any],
    extraction_table: List[Dict[str, Any]],
    api_key: str = None,
    **kwargs
) -> str:
    """
    Generate full research report and return as markdown.

    Args:
        research_question: Main research question
        scopus_metadata: Scopus search statistics
        extraction_table: Data extraction results
        api_key: Anthropic API key
        **kwargs: Additional parameters for generate_full_report

    Returns:
        Complete report in markdown format
    """
    orchestrator = NarrativeOrchestrator(api_key=api_key)
    orchestrator.generate_full_report(
        research_question=research_question,
        scopus_metadata=scopus_metadata,
        extraction_table=extraction_table,
        **kwargs
    )
    return orchestrator.export_to_markdown()
